import streamlit as st
import boto3
from io import BytesIO
from datetime import datetime
import requests
import re
import textwrap

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

DIAGRAM_TYPES = [
    "flowchart",
    "sequence",
    "class",
    "state",
    "entity_relationship",
    "gantt",
    "pie",
    "gitgraph",
]

def set_background():
    st.markdown("""
        <style>
        body, .stApp {background: #f6fbff;}
        .stTextArea textarea, .stFileUploader {background: #f7fbff !important;}
        .stButton button {background-color: #1976d2 !important; color: #fff !important; border-radius: 6px !important;}
        .stDownloadButton, .stDownloadButton button {background: #2196f3 !important; color: #fff !important; border-radius: 6px !important;}
        .rfc-preview-area {background: #e3f2fd; border: 1.5px solid #2196F3; border-radius: 10px; padding: 1.5em; font-size: 1em; overflow-x: auto; margin-bottom: 0.7em;}
        .stepper {margin-bottom: 15px; font-weight: bold; color: #1976d2; background: #e3f2fd; padding: 7px 18px; border-radius: 8px;}
        </style>
    """, unsafe_allow_html=True)

def show_stepper(current):
    steps = [
        "Login/Select RFC",
        "File Selection",
        "Metadata",
        "Diagram Conversion",
        "RFC Review & Export",
        "Comment on RFC"
    ]
    txt = "<div class='stepper'>"
    for idx, step in enumerate(steps):
        color = "#1976d2" if idx == current else "#bbb"
        txt += f"<span style='margin-right:10px; color:{color};'>{step}{' ➔' if idx < len(steps)-1 else ''}</span>"
    txt += "</div>"
    st.markdown(txt, unsafe_allow_html=True)

def get_s3_client(aws_access_key, aws_secret_key, region):
    return boto3.client(
        "s3",
        region_name=region,
        aws_access_key_id=aws_access_key,
        aws_secret_access_key=aws_secret_key,
    )

def list_buckets(s3_client):
    try:
        return [bucket["Name"] for bucket in s3_client.list_buckets().get("Buckets", [])]
    except Exception:
        return []

def list_objects(s3_client, bucket):
    try:
        return [obj["Key"] for obj in s3_client.list_objects_v2(Bucket=bucket).get("Contents", [])]
    except Exception:
        return []

def upload_to_s3(s3_client, bucket, filename, content_bytes):
    s3_client.upload_fileobj(BytesIO(content_bytes), bucket, filename)

def get_bedrock_agent_client(region, aws_access_key, aws_secret_key):
    return boto3.client(
        service_name="bedrock-agent-runtime",
        region_name=region,
        aws_access_key_id=aws_access_key,
        aws_secret_access_key=aws_secret_key,
    )

def bedrock_agent_ask(agent_id, alias_id, region, aws_access_key, aws_secret_key, user_message):
    try:
        client = get_bedrock_agent_client(region, aws_access_key, aws_secret_key)
        response_stream = client.invoke_agent(
            agentId=agent_id,
            agentAliasId=alias_id,
            sessionId="rfc-session",
            inputText=user_message,
        )
        output = ""
        for event in response_stream['completion']:
            if "chunk" in event and "bytes" in event["chunk"]:
                output += event["chunk"]["bytes"].decode("utf-8")
        return output.strip()
    except Exception as e:
        return f"[Bedrock Agent Error: {e}]"

def extract_images_from_docx(file_bytes):
    images = []
    if Document is None:
        return images
    doc = Document(BytesIO(file_bytes))
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img_blob = rel.target_part.blob
            images.append(img_blob)
    return images

def extract_images_from_pdf(file_bytes):
    images = []
    if PyPDF2 is None:
        return images
    try:
        import fitz
    except ImportError:
        return images
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    for page in doc:
        for img in page.get_images(full=True):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            images.append(image_bytes)
    return images

def extract_text_from_file(file_bytes, file_name):
    ext = file_name.lower().split(".")[-1]
    if ext in ["md", "txt"]:
        return file_bytes.decode("utf-8", errors="ignore")
    elif ext == "docx" and Document is not None:
        try:
            doc = Document(BytesIO(file_bytes))
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception:
            return "[Could not extract text from DOCX]"
    elif ext == "pdf" and PyPDF2 is not None:
        try:
            reader = PyPDF2.PdfReader(BytesIO(file_bytes))
            return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
        except Exception:
            return "[Could not extract text from PDF]"
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return "[Unsupported file type or unreadable content]"

def extract_images_from_file(file_bytes, file_name):
    ext = file_name.lower().split(".")[-1]
    if ext == "docx":
        return extract_images_from_docx(file_bytes)
    elif ext == "pdf":
        return extract_images_from_pdf(file_bytes)
    elif ext in ["png", "jpg", "jpeg"]:
        return [file_bytes]
    else:
        return []

def ocr_image_bytes(image_bytes):
    if Image is None or pytesseract is None:
        return "[OCR not available: Tesseract or Pillow is not installed]"
    try:
        img = Image.open(BytesIO(image_bytes))
        return pytesseract.image_to_string(img)
    except Exception as e:
        return f"[OCR not available: {e}]"

def extract_metadata_from_markdown(md_text):
    regex = r"\*\*Author:\*\* (.*?)\s*[\r\n]+.*?\*\*Date:\*\* (.*?)\s*[\r\n]+.*?\*\*Status:\*\* (.*?)\s*[\r\n]+.*?\*\*Reviewers:\*\* (.*?)\s*[\r\n]+.*?\*\*Topic:\*\* (.*?)\s*[\r\n]+"
    m = re.search(regex, md_text, re.DOTALL)
    if m:
        return {
            "author": m.group(1).strip(),
            "date": m.group(2).strip(),
            "status": m.group(3).strip(),
            "reviewers": m.group(4).strip(),
            "topic": m.group(5).strip()
        }
    return {}

def update_metadata_in_markdown(md_text, new_metadata):
    pattern = r'(\*\*Author:\*\*\s*)(.*?)(\s*[\r\n]+)(.*?)(\*\*Date:\*\*\s*)(.*?)(\s*[\r\n]+)(.*?)(\*\*Status:\*\*\s*)(.*?)(\s*[\r\n]+)(.*?)(\*\*Reviewers:\*\*\s*)(.*?)(\s*[\r\n]+)(.*?)(\*\*Topic:\*\*\s*)(.*?)(\s*[\r\n]+)'
    repl = (
        r'\g<1>{author}\g<3>\g<4>\g<5>{date}\g<7>\g<8>\g<9>{status}\g<11>\g<12>\g<13>{reviewers}\g<15>\g<16>\g<17>{topic}\g<19>'
    ).format(
        author=new_metadata.get("author", ""),
        date=new_metadata.get("date", ""),
        status=new_metadata.get("status", ""),
        reviewers=new_metadata.get("reviewers", ""),
        topic=new_metadata.get("topic", "")
    )
    new_md = re.sub(pattern, repl, md_text, flags=re.DOTALL)
    return new_md

def extract_external_diagram_links(text):
    image_urls = re.findall(r'(https?://[^\s]+?\.(?:png|jpg|jpeg|gif))', text)
    drawio_urls = re.findall(r'(https?://[^\s]*draw\.io[^\s]*)', text)
    mermaid_urls = re.findall(r'(https?://mermaid\.live[^\s]*)', text)
    mermaid_blocks = re.findall(r'```mermaid\n(.*?)\n```', text, re.DOTALL)
    return image_urls, drawio_urls, mermaid_urls, mermaid_blocks

def download_image(url):
    try:
        resp = requests.get(url)
        if resp.ok:
            return resp.content
    except Exception:
        return None
    return None

def extract_mermaid_from_drawio(drawio_url):
    return "[draw.io to mermaid conversion not implemented]"

def extract_mermaid_from_mermaid_url(mermaid_url):
    try:
        import urllib.parse
        u = urllib.parse.urlparse(mermaid_url)
        code = ""
        if u.fragment:
            code = urllib.parse.unquote(u.fragment)
        else:
            qs = urllib.parse.parse_qs(u.query)
            text = qs.get('graph', [""])[0]
            code = urllib.parse.unquote(text)
        return code
    except Exception:
        return "[Could not extract mermaid code from URL]"

def process_document_for_diagrams(file_bytes, file_name, file_text=None):
    images = extract_images_from_file(file_bytes, file_name)
    if file_text is None:
        file_text = extract_text_from_file(file_bytes, file_name)
    image_urls, drawio_urls, mermaid_urls, mermaid_blocks = extract_external_diagram_links(file_text)
    for url in image_urls:
        img_bytes = download_image(url)
        if img_bytes:
            images.append(img_bytes)
    mermaid_diagrams = []
    for drawio_url in drawio_urls:
        mermaid_diagrams.append(extract_mermaid_from_drawio(drawio_url))
    for mermaid_url in mermaid_urls:
        mermaid_diagrams.append(extract_mermaid_from_mermaid_url(mermaid_url))
    for mb in mermaid_blocks:
        mermaid_diagrams.append(mb.strip())
    return images, mermaid_diagrams

def extract_mermaid_code(response):
    m = re.search(r"```mermaid(.*?)(```|$)", response, re.DOTALL | re.IGNORECASE)
    if m:
        return m.group(1).strip()
    return response.strip()

def mermaid_from_image(description, session_state, diagram_type):
    if not description.strip():
        return "[No OCR or image description available]"
    prompt = (
        f"Given the following diagram description, generate a detailed Mermaid diagram using type '{diagram_type}'. "
        f"Return ONLY the Mermaid code block (inside ```mermaid ... ```), and NOTHING else.\n"
        f"Diagram description:\n{description}\n"
    )
    response = bedrock_agent_ask(
        session_state['bedrock_agent_id'],
        session_state['bedrock_alias_id'],
        session_state['bedrock_region'],
        session_state['aws_access_key'],
        session_state['aws_secret_key'],
        prompt
    )
    return extract_mermaid_code(response)

def markdown_from_ai(doc_text, image_mermaids, metadata, session_state, diagram_type, custom_prompt=None, prompt_image_text=None):
    author = metadata.get("author") or "Unknown"
    topic = metadata.get("topic") or "Unknown"
    date = metadata.get("date") or datetime.now().strftime("%Y-%m-%d")
    status = metadata.get("status") or "Draft"
    reviewers = metadata.get("reviewers") or "N/A"
    toc = (
        "1. [Abstract](#1-abstract)\n"
        "2. [Status of This Memo](#2-status-of-this-memo)\n"
        "3. [Introduction](#3-introduction)\n"
        "4. [Motivation](#4-motivation)\n"
        "5. [Proposal](#5-proposal)\n"
        "6. [Architecture](#6-architecture)\n"
        "7. [Diagrams](#7-diagrams)\n"
        "8. [Alternatives](#8-alternatives)\n"
        "9. [Risks](#9-risks)\n"
        "10. [Implementation Plan](#10-implementation-plan)\n"
        "11. [Security Considerations](#11-security-considerations)\n"
        "12. [References](#12-references)\n"
        "13. [Acknowledgements](#13-acknowledgements)\n"
    )
    diagrams = ""
    for idx, mermaid in enumerate(image_mermaids):
        diagrams += f"\n#### Diagram {idx+1}\n```mermaid\n{mermaid}\n```\n"
    prompt = f"""
You are an RFC Markdown generator.
Given the extracted document text below, output a complete RFC document in markdown format ONLY.
The document MUST use this structure:

# RFC: {topic}

**Author:** {author}  
**Date:** {date}  
**Status:** {status}  
**Reviewers:** {reviewers}  
**Topic:** {topic}

---

## Table of Contents

{toc}

---

## 1. Abstract

---

## 2. Status of This Memo

---

## 3. Introduction

---

## 4. Motivation

---

## 5. Proposal

---

## 6. Architecture

---

## 7. Diagrams

{diagrams}

---

## 8. Alternatives

---

## 9. Risks

---

## 10. Implementation Plan

---

## 11. Security Considerations

---

## 12. References

---

## 13. Acknowledgements

---

Extracted document text:
{textwrap.shorten(doc_text, width=5000, placeholder='...')}

- Fill each section with content from the provided text if possible. Use placeholder sentences if necessary.
- When rendering diagrams, always use fenced Markdown code blocks with the mermaid language identifier.
- Do not summarize, explain, or provide any output other than the RFC Markdown.
- Never provide commentary outside the Markdown.
"""
    if custom_prompt:
        prompt += f"\n\nAdditional user prompt: {custom_prompt}"
    if prompt_image_text:
        prompt += f"\n\nExtra context from uploaded file/image:\n{prompt_image_text}\n"
    return bedrock_agent_ask(
        session_state['bedrock_agent_id'],
        session_state['bedrock_alias_id'],
        session_state['bedrock_region'],
        session_state['aws_access_key'],
        session_state['aws_secret_key'],
        prompt
    )

def login_ui():
    set_background()
    show_stepper(0)
    st.markdown("<h1 style='color:#2A60C0'> RFC AI Agent </h1>", unsafe_allow_html=True)
    with st.form("login_form"):
        aws_access_key = st.text_input("AWS Access Key ID", type="password")
        aws_secret_key = st.text_input("AWS Secret Access Key", type="password")
        bedrock_agent_id = st.text_input("Bedrock Agent ID")
        bedrock_alias_id = st.text_input("Bedrock Alias ID")
        bedrock_region = st.text_input("Bedrock Region", value="us-east-1")
        submitted = st.form_submit_button("Login")
    if submitted:
        if not (aws_access_key and aws_secret_key and bedrock_agent_id and bedrock_alias_id and bedrock_region):
            st.error("Fill in all fields.")
            return
        try:
            s3_client = get_s3_client(aws_access_key, aws_secret_key, bedrock_region)
            st.session_state.s3_client = s3_client
            st.session_state.aws_access_key = aws_access_key
            st.session_state.aws_secret_key = aws_secret_key
            st.session_state.bedrock_agent_id = bedrock_agent_id
            st.session_state.bedrock_alias_id = bedrock_alias_id
            st.session_state.bedrock_region = bedrock_region
            st.session_state.stage = "choose_rfc_or_new"
            st.rerun()
        except Exception as e:
            st.error(f"Login failed: {e}")

def choose_rfc_or_new_ui():
    set_background()
    show_stepper(0)
    st.markdown("<h2 style='color:#1976D2'>Choose an RFC Operation</h2>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📝 New RFC Document", key="choose_new_rfc", use_container_width=True):
            st.session_state.choose_rfc_flow = "new"
            st.session_state.stage = "file_choice"
            st.rerun()
    with col2:
        if st.button("Review/Comment on Existing RFC", key="choose_review_rfc", use_container_width=True):
            st.session_state.choose_rfc_flow = "existing"
            st.session_state.stage = "manager_comment"
            st.rerun()
    back = st.button("⬅️ Back", key="choose_rfc_back")
    if back:
        st.session_state.stage = "login"
        st.rerun()

def file_choice_ui():
    set_background()
    show_stepper(1)
    st.markdown("<h2 style='color:#1976D2'>Step 1: Select or Upload a File</h2>", unsafe_allow_html=True)
    st.markdown('<div class="file-choice-div">', unsafe_allow_html=True)
    col_upload, col_s3 = st.columns(2)
    with col_upload:
        st.subheader("Upload from PC")
        uploaded_file = st.file_uploader(
            "Upload a Word, PDF, Markdown, or Text file:",
            type=["docx", "pdf", "md", "txt"],
            key="upload_any"
        )
        if uploaded_file:
            file_content = uploaded_file.read()
            file_name = uploaded_file.name
            st.session_state.file_origin = "upload"
            st.session_state.file_content = file_content
            st.session_state.file_name = file_name
            st.success(f"File '{file_name}' uploaded.")
    with col_s3:
        st.subheader("Select from S3 Bucket")
        if "s3_client" in st.session_state:
            buckets = list_buckets(st.session_state.s3_client)
            bucket = st.selectbox("S3 Bucket", buckets, key="bucket_select")
            if bucket:
                objects = list_objects(st.session_state.s3_client, bucket)
                if objects:
                    s3_file = st.selectbox("File in bucket", objects, key="s3_file_select")
                    if st.button("Load file from S3", key="load_s3_file"):
                        file_obj = BytesIO()
                        st.session_state.s3_client.download_fileobj(bucket, s3_file, file_obj)
                        file_obj.seek(0)
                        file_content = file_obj.read()
                        file_name = s3_file
                        st.session_state.file_origin = "s3"
                        st.session_state.file_content = file_content
                        st.session_state.file_name = file_name
                        st.session_state.bucket = bucket
                        st.success(f"Loaded '{file_name}' from S3. Click Continue to proceed.")
        else:
            st.info("Please login to list S3 buckets.")
    st.markdown('</div>', unsafe_allow_html=True)
    col_btn = st.columns([1, 1, 1])
    back = col_btn[0].button("⬅️ Back", key="back_to_choose_rfc", use_container_width=True)
    continue_btn = col_btn[1].button(
        "Continue",
        key="continue_file_choice",
        use_container_width=True,
        disabled=not (st.session_state.get("file_content") and st.session_state.get("file_name"))
    )
    if continue_btn:
        st.session_state.stage = "metadata"
        st.rerun()
    if back:
        st.session_state.stage = "choose_rfc_or_new"
        st.rerun()

def metadata_ui():
    set_background()
    show_stepper(2)
    st.markdown("<h2 style='color:#1976D2'>Step 2: Enter RFC Metadata</h2>", unsafe_allow_html=True)
    if "metadata" not in st.session_state:
        st.session_state.metadata = {
            "author": "",
            "topic": "",
            "status": "Draft",
            "reviewers": "",
            "date": datetime.now().strftime("%Y-%m-%d"),
        }
    metadata = st.session_state.metadata
    author = st.text_input("Author", value=metadata.get("author", ""))
    topic = st.text_input("Topic", value=metadata.get("topic", ""))
    status = st.selectbox("Status", ["Draft", "Accepted", "Rejected", "Implemented"], index=["Draft", "Accepted", "Rejected", "Implemented"].index(metadata.get("status", "Draft")))
    reviewers = st.text_input("Reviewers (comma-separated)", value=metadata.get("reviewers", ""))
    date = st.text_input("Date", value=metadata.get("date", datetime.now().strftime("%Y-%m-%d")))
    if st.button("Extract & Convert with AI"):
        images, mermaid_diagrams = process_document_for_diagrams(st.session_state.file_content, st.session_state.file_name)
        text = extract_text_from_file(st.session_state.file_content, st.session_state.file_name)
        st.session_state.metadata = {
            "author": author, "topic": topic, "status": status,
            "reviewers": reviewers, "date": date,
        }
        st.session_state.text_content = text
        st.session_state.images = images
        st.session_state.external_mermaid_diagrams = mermaid_diagrams
        st.session_state.stage = "image_mermaid"
        st.rerun()
    if st.button("⬅️ Back to File Selection", key="back_to_filechoice", help="Back to file selection", type="primary"):
        st.session_state.stage = "file_choice"
        st.rerun()

def image_mermaid_ui():
    set_background()
    show_stepper(3)
    st.markdown("<h2 style='color:#1976D2'>Step 3: Convert Images to Mermaid Diagrams</h2>", unsafe_allow_html=True)
    st.write("Extracted images and existing diagrams from the file are shown below. For each image, review/correct the OCR text or describe the diagram, then generate and edit its Mermaid code as needed. Existing Mermaid diagrams are also shown for editing.")
    images = st.session_state.images
    external_mermaid_diagrams = st.session_state.external_mermaid_diagrams
    image_mermaids = []
    st.markdown("**Diagram type for Mermaid generation:**")
    if "diagram_type" not in st.session_state:
        st.session_state.diagram_type = "flowchart"
    st.session_state.diagram_type = st.selectbox(
        "Choose Mermaid diagram type",
        options=DIAGRAM_TYPES,
        key="diagram_type_select",
    )
    chosen_type = st.session_state.diagram_type
    if len(images) == 0 and len(external_mermaid_diagrams) == 0:
        st.info("No images or diagrams detected in the file. You can proceed to RFC draft directly.")
    for idx, img_bytes in enumerate(images):
        st.image(img_bytes, caption=f"Image {idx+1}", use_container_width=True)
        ocr_text = ocr_image_bytes(img_bytes)
        ocr_text_key = f"ocr_text_{idx}"
        st.markdown(f"**OCR for Image {idx+1}** *(edit or describe if needed)*:")
        ocr_text = st.text_area(f"OCR Text for Image {idx+1}", value=ocr_text, key=ocr_text_key)
        if not ocr_text.strip():
            ocr_text = st.text_area(
                f"Describe the diagram for Image {idx+1} (if OCR failed):", key=f"manual_{ocr_text_key}"
            )
        key = f"mermaid_{idx}"
        if st.button(f"Generate Mermaid for Image {idx+1}", key=f"btn_mermaid_{idx}"):
            with st.spinner(f"Generating Mermaid diagram for Image {idx+1}..."):
                mermaid = mermaid_from_image(
                    ocr_text, st.session_state, chosen_type
                )
                st.session_state[key] = mermaid
        mermaid_code = st.text_area(
            f"Mermaid for Image {idx+1} (edit as needed)",
            value=st.session_state.get(key, ""),
            height=180,
            key=f"mermaid_code_{idx}",
        )
        st.session_state[key] = mermaid_code
        image_mermaids.append(mermaid_code)
        st.markdown("#### Diagram Preview:")
        st.markdown(f"```mermaid\n{mermaid_code}\n```")
    for idx, mermaid_code in enumerate(external_mermaid_diagrams):
        key = f"external_mermaid_{idx}"
        code = st.text_area(
            f"External Mermaid Diagram {idx+1} (edit as needed)",
            value=mermaid_code,
            height=180,
            key=f"external_mermaid_code_{idx}",
        )
        image_mermaids.append(code)
        st.markdown("#### Diagram Preview:")
        st.markdown(f"```mermaid\n{code}\n```")
    st.session_state.image_mermaids = image_mermaids
    if st.button("Generate RFC Markdown with AI"):
        with st.spinner("Generating RFC Markdown with AI..."):
            markdown = markdown_from_ai(
                st.session_state.text_content,
                st.session_state.image_mermaids,
                st.session_state.metadata,
                st.session_state,
                chosen_type,
            )
        st.session_state.md_code = markdown
        st.session_state.md_code_edit = markdown
        st.session_state.stage = "md_review"
        st.rerun()
    if st.button("⬅️ Back to Metadata", key="back_to_metadata", help="Back to metadata step", type="primary"):
        st.session_state.stage = "metadata"
        st.rerun()

def md_review_ui():
    set_background()
    show_stepper(4)
    st.markdown("<h2 style='color:#1976D2'>Step 4: Review, Edit & Export RFC Markdown</h2>", unsafe_allow_html=True)
    tab_preview, tab_code = st.tabs(["Preview", "Edit Markdown"])
    if "md_code_edit" not in st.session_state:
        st.session_state.md_code_edit = st.session_state.md_code
    with tab_code:
        new_code = st.text_area(
            "Edit Markdown",
            value=st.session_state.md_code_edit,
            height=600,
            key="edit_md",
        )
        if new_code != st.session_state.md_code_edit:
            st.session_state.md_code_edit = new_code
    with tab_preview:
        st.markdown('<div class="rfc-preview-area">', unsafe_allow_html=True)
        st.markdown(st.session_state.md_code_edit, unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("<h3 style='color:#1565C0;'>AI-Driven RFC Update & Regeneration</h3>", unsafe_allow_html=True)
    with st.form("ai_rfc_form"):
        user_prompt = st.text_area(
            "Prompt for AI (describe changes, add requirements, ask for improvements, etc.):",
            value="",
            key="ai_user_prompt",
            height=200,
            placeholder="Type your RFC review prompt here. You can copy-paste text or describe image/file content."
        )
        prompt_file = st.file_uploader(
            "Upload file/image for prompt context (optional)",
            type=["png","jpg","jpeg","txt","md","pdf","docx"],
            key="prompt_file"
        )
        chosen_type = st.selectbox(
            "Choose Mermaid Diagram Type (for regeneration)",
            DIAGRAM_TYPES,
            index=DIAGRAM_TYPES.index(st.session_state.diagram_type),
            key="regen_mermaid_diagram_type"
        )
        regenerate = st.form_submit_button("Regenerate RFC Document with AI")
    prompt_image_text = ""
    if prompt_file is not None:
        prompt_file_bytes = prompt_file.read()
        prompt_file_name = prompt_file.name
        if prompt_file.type.startswith('image/'):
            if Image is not None:
                prompt_image_text = ocr_image_bytes(prompt_file_bytes)
        else:
            prompt_image_text = extract_text_from_file(prompt_file_bytes, prompt_file_name)
        st.success(f"Context file '{prompt_file_name}' attached.")
    if regenerate:
        with st.spinner("Updating RFC document with AI..."):
            new_image_mermaids = []
            for idx, img_bytes in enumerate(st.session_state.images):
                ocr_desc = ocr_image_bytes(img_bytes)
                if not ocr_desc.strip():
                    ocr_desc = st.session_state.get(f"ocr_text_{idx}", "")
                mermaid = mermaid_from_image(
                    ocr_desc,
                    st.session_state,
                    chosen_type
                )
                new_image_mermaids.append(mermaid)
            for idx, mermaid_code in enumerate(st.session_state.external_mermaid_diagrams):
                key = f"external_mermaid_{idx}"
                code = st.session_state.get(key, mermaid_code)
                new_image_mermaids.append(code)
            st.session_state.image_mermaids = new_image_mermaids
            markdown = markdown_from_ai(
                st.session_state.md_code_edit,
                st.session_state.image_mermaids,
                st.session_state.metadata,
                st.session_state,
                chosen_type,
                custom_prompt=user_prompt,
                prompt_image_text=prompt_image_text
            )
            st.session_state.md_code = markdown
            st.session_state.md_code_edit = markdown
            st.session_state.diagram_type = chosen_type
            st.rerun()
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Upload RFC Markdown Document to S3"):
            upload_bucket = st.session_state.get("bucket") or st.selectbox(
                "Select S3 bucket to upload .md",
                list_buckets(st.session_state.s3_client),
                key="up_bucket",
            )
            md_filename = st.text_input(
                "File name for RFC Markdown",
                value=st.session_state.file_name.rsplit(".", 1)[0] + "_RFC.md",
            )
            if upload_bucket and md_filename:
                try:
                    upload_to_s3(
                        st.session_state.s3_client,
                        upload_bucket,
                        md_filename,
                        st.session_state.md_code_edit.encode("utf-8"),
                    )
                    st.success(
                        f"RFC Markdown file uploaded to '{upload_bucket}' as '{md_filename}'!"
                    )
                except Exception as e:
                    st.error(f"Upload failed: {e}")
    with col2:
        st.download_button(
            label="Download RFC Markdown Document",
            data=st.session_state.md_code_edit,
            file_name=st.session_state.file_name.rsplit(".", 1)[0] + "_RFC.md",
            mime="text/markdown",
        )
    st.markdown("---")
    if st.button("⬅️ Back to Diagram Conversion", key="back_to_image_mermaid", help="Back to diagram conversion step", type="primary"):
        st.session_state.stage = "image_mermaid"
        st.rerun()
    if st.button("🔁 Start Over (Logout)", key="logout_rfc", help="Logout and start over"):
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.session_state.stage = "login"
        st.rerun()

def manager_comment_ui():
    set_background()
    show_stepper(5)
    st.markdown("<h2 style='color:#1976D2'>Step: Review & Comment on Existing RFC</h2>", unsafe_allow_html=True)
    st.markdown("Select a bucket and an existing RFC Markdown file from S3 to review and comment on:")
    buckets = list_buckets(st.session_state.s3_client)
    bucket = st.selectbox("Select a bucket", buckets, key="mgr_bucket_select")
    rfc_file = None
    rfc_md = ""
    if bucket:
        objects = [f for f in list_objects(st.session_state.s3_client, bucket) if f.lower().endswith(".md")]
        rfc_file = st.selectbox("Select an RFC Markdown file", objects, key="mgr_file_select")
        if rfc_file:
            file_obj = BytesIO()
            st.session_state.s3_client.download_fileobj(bucket, rfc_file, file_obj)
            file_obj.seek(0)
            rfc_md = file_obj.read().decode("utf-8", errors="ignore")
    if rfc_md:
        st.markdown("#### RFC Metadata (Editable)")
        md_metadata = extract_metadata_from_markdown(rfc_md)
        if not md_metadata:
            st.info("No RFC metadata fields detected. The file may not be in standard RFC format.")
        else:
            col1, col2 = st.columns(2)
            with col1:
                author = st.text_input("Author", value=md_metadata.get("author", ""))
                status = st.selectbox("Status", ["Draft", "Accepted", "Rejected", "Implemented"], index=["Draft", "Accepted", "Rejected", "Implemented"].index(md_metadata.get("status", "Draft")))
            with col2:
                reviewers = st.text_input("Reviewers (comma-separated)", value=md_metadata.get("reviewers", ""))
                date = st.text_input("Date", value=md_metadata.get("date", ""))
            topic = st.text_input("Topic", value=md_metadata.get("topic", ""))
            updated_metadata = dict(author=author, status=status, reviewers=reviewers, date=date, topic=topic)
            updated_md = update_metadata_in_markdown(rfc_md, updated_metadata)
            st.markdown('<div class="rfc-preview-area">', unsafe_allow_html=True)
            st.markdown(updated_md)
            st.markdown('</div>', unsafe_allow_html=True)
        comment = st.text_area("Manager Comments", placeholder="Write your comments or review feedback here...")
        append_comments = st.checkbox("Append comments to RFC Markdown (at the end)", value=True)
        final_md = updated_md if md_metadata else rfc_md
        if comment.strip() and append_comments:
            final_md += f"\n\n---\n## Manager Comments\n{comment}\n"
        elif comment.strip():
            final_md = f"## Manager Comments\n{comment}\n\n---\n" + final_md
        st.markdown("#### Updated RFC Markdown (with comments)")
        st.text_area("RFC Markdown (with comments)", value=final_md, height=350, key="mgr_md_edit")
        st.download_button(
            label="Download RFC Markdown (with comments)",
            data=final_md,
            file_name=(rfc_file.rsplit(".", 1)[0] + "_with_comments.md") if rfc_file else "rfc_with_comments.md",
            mime="text/markdown",
        )
        if st.button("Upload Updated RFC with Comments to S3"):
            upload_filename = (rfc_file.rsplit(".", 1)[0] + "_with_comments.md") if rfc_file else "rfc_with_comments.md"
            try:
                upload_to_s3(st.session_state.s3_client, bucket, upload_filename, final_md.encode("utf-8"))
                st.success(f"RFC Markdown with comments uploaded as '{upload_filename}' to bucket '{bucket}'!")
            except Exception as e:
                st.error(f"Upload failed: {e}")
    if st.button("⬅️ Back", key="back_to_choose", help="Back to option selection", type="primary"):
        st.session_state.stage = "choose_rfc_or_new"
        st.rerun()

if "stage" not in st.session_state:
    st.session_state.stage = "login"
if st.session_state.stage == "login":
    login_ui()
elif st.session_state.stage == "choose_rfc_or_new":
    choose_rfc_or_new_ui()
elif st.session_state.stage == "manager_comment":
    manager_comment_ui()
elif st.session_state.stage == "file_choice":
    file_choice_ui()
elif st.session_state.stage == "metadata":
    metadata_ui()
elif st.session_state.stage == "image_mermaid":
    image_mermaid_ui()
elif st.session_state.stage == "md_review":
    md_review_ui()